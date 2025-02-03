import os
import zlib
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
#import comtypes.client
import pdfplumber
import traceback

def get_file_info(file_path):
    """Получаем имя, размер, дату изменения и хеш-сумму файла (CRC32)"""
    file_name = os.path.basename(file_path)  # Имя файла
    
    size = os.path.getsize(file_path)  # Размер в байтах
    modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')  # Дата изменения
    
    # Хеш-сумма CRC32
    crc32_hash = 0
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            crc32_hash = zlib.crc32(chunk, crc32_hash)
    file_hash = format(crc32_hash & 0xFFFFFFFF, '08X')

 
    region = (70, 400, 565, 590)
    text = extract_text_pymupdf(file_path, region)
    
    return file_name, size, modified_time, file_hash, text

def update_template(template_path, output_path, file_info):
    """Обновляет только текст в ячейках таблицы в документе, сохраняет форматирование"""
    doc = Document(template_path)
    replacements = {
        '{FILE_NAME}': file_info[0],
        '{SIZE}': str(file_info[1]),
        '{DATE}': file_info[2],
        '{HASH}': file_info[3],
        '{TEXT}': file_info[4]
    }
    
    # Заменяем текст только в ячейках таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        # Проходим по всем параграфам в ячейке
                        for paragraph in cell.paragraphs:
                            # Проходим по каждому run в параграфе
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
    
    doc.save(output_path)



import sys

if sys.platform == "win32":
    import comtypes.client

    def convert_docx_to_pdf(docx_path, pdf_path):
        word = comtypes.client.CreateObject("Word.Application")
        print(docx_path)
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
else:
    import subprocess
    def convert_docx_to_pdf(docx_path, pdf_path):
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", pdf_path.rsplit("/", 1)[0], docx_path])



def extract_text_pymupdf(pdf_path, bbox):
    """
    Извлекает текст из заданного региона первой страницы PDF.
    
    :param pdf_path: путь к PDF-файлу
    :param bbox: кортеж (x0, top, x1, bottom) - координаты региона
    :return: извлеченный текст
    """
    with pdfplumber.open(pdf_path) as pdf:
        first_page = pdf.pages[0]  # Первая страница
        #print(first_page.extract_words())  # Покажет координаты всех слов
        text = first_page.within_bbox(bbox).extract_text()  # Извлекаем текст из региона
        return text




def main():
    input_folder = 'pdf_files'
    output_folder = 'output'
    template_file = 'template.docx'
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    for file_name in os.listdir(input_folder):
        if file_name.lower().endswith('.pdf'):
            file_path = os.path.join(input_folder, file_name)
           
            print(f'Обрабатываем файл: {file_name}')
            file_info = get_file_info(file_path)


            # Создаем путь для временного DOCX файла
            temp_docx = os.path.abspath(os.path.join(output_folder, f'{os.path.splitext(file_name)[0]}.docx'))
            
            # Создаем путь для финального PDF файла, убираем расширение '.pdf'
            final_pdf = os.path.abspath(os.path.join(output_folder, os.path.splitext(file_name)[0]))
            
            update_template(template_file, temp_docx, file_info)
            convert_docx_to_pdf(temp_docx, final_pdf)
            os.remove(temp_docx)  # Удаляем временный DOCX
            print(f'Создан файл: {final_pdf}')

try:
    main()
except Exception as e:
    print("Произошла ошибка:")
    print(e)
    print(traceback.format_exc())	
    input("Нажмите Enter для выхода...")